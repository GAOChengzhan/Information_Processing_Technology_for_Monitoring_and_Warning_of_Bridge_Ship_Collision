# -*- coding: utf-8 -*-
"""
Created on Wed Jul 18 22:16:31 2018

@author: DELL
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import xlrd
#文字
path=r"C:\Users\DELL\Desktop\1.docx"
document = Document()
pic = document.add_picture('4.jpg',width=Inches(4))
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
#ocument.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
document.add_heading('自动生成报告测试', 0)
document.add_heading('Test For Automatically Generate Reports', level=1)
document.add_page_break()#增加分页
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True
document.add_heading('This is my title', 0)
document.add_paragraph('my paragraph')
document.styles['Normal'].font.name = u'黑体'
p = document.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER#居中
run = p.add_run(u'我添加的段落文字 ')
run.font.color.rgb = RGBColor(54, 95, 145)#颜色
run.font.size = Pt(39)#字号
#测试缩进
paragraph = document.add_paragraph('测试左侧缩进0.3英寸')
paragraph_format = paragraph.paragraph_format
paragraph_format.left_indent = Inches(0.3)
#首行缩进
paragraph = document.add_paragraph('测试首行缩进1英寸')
paragraph_format1 = paragraph.paragraph_format
paragraph_format1.first_line_indent = Inches(1)
#段落测试
paragraph1=[]
d=open("C:\\Users\\DELL\\Desktop\\自动报告尝试1\\docsample\\测试.txt","r",True)
for line in d:
    temp=line[:-1]
    document.add_paragraph(temp)
    paragraph1.append(temp)
paragraph=str(paragraph1)
paragraph_format.space_before = Pt(1)#本段文字与上一段的间距
paragraph_format.space_after = Pt(24)#本段文字与下一段的间距 
paragraph_format.line_spacing = Pt(18)#行间距
#分页测试
paragraph_format.keep_together#紧跟上段
paragraph_format.keep_with_next#若本页无法完全显示，另起一页
#字体格式： 
p = document.add_paragraph()
run = p.add_run('字体测试')
#加粗
run.font.bold = True
#斜体
run.font.italic = True
#下划线
run.font.underline = True
#颜色
run.font.color.rgb = RGBColor(0x42, 0x24 , 0xE9)
#用txt做模拟
for i in range(1,4):    
    e=open("C:\\Users\\DELL\\Desktop\\自动报告尝试1\\docsample\\"+str(i)+".txt","r",True)
    for line in e:
        temp=line[:-1]
        document.add_paragraph(temp)    
#图片
pic = document.add_picture('1.jpg')
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
document.add_picture('2.jpg', width=Inches(4))#调节大小
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
document.add_picture('3.jpg', width=Inches(2.5))
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#表格
rows = 2
cols = 3
table = document.add_table(rows=rows, cols=cols,style = "Table Grid")  # 添加2行3列的表格
 
for i in range(rows):
    tr = table.rows[i]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), "450")
    trPr.append(trHeight)  # 表格高度设置
# table.autofit = False
col = table.columns[1]
col.width = Inches(5)
arr = [u'序号',u"类型",u"详细描述"]
heading_cells = table.rows[0].cells
for i in range(cols):
    p = heading_cells[i].paragraphs[0]
    run = p.add_run(arr[i])
    run.font.color.rgb = RGBColor(54, 95, 145)  # 颜色设置，这里是用RGB颜色
    run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
table.cell(1, 1).text = u'表格文字'
table.add_row()
#用excel模拟
x=xlrd.open_workbook(r"C:\Users\DELL\Desktop\自动报告尝试1\docsample\1.xls")
y=x.sheets()[0]#获得第一个sheet
excelh=y.nrows
excell=y.ncols
data=[]
for i in range(0,y.nrows):
    a=y.row_values(i)
    data.append(a)
table = document.add_table(rows=excelh, cols=excell,style = "Table Grid")
for i in range(excelh):
    tr = table.rows[i]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), "450")
    trPr.append(trHeight)  # 表格高度设置        
for j in range(excelh):
    array=data[j-1]
    for i in range(excell):
        x_cells = table.rows[j-1].cells
        p = x_cells[i].paragraphs[0]
        run = p.add_run(array[i])
#换一个再测试一次
x=xlrd.open_workbook(r"C:\Users\DELL\Desktop\自动报告尝试1\docsample\!.xlsx")
y=x.sheets()[0]#获得第一个sheet
excelh=y.nrows
excell=y.ncols
data=[]
for i in range(0,y.nrows):
    a=y.row_values(i)
    data.append(a)
table = document.add_table(rows=excelh, cols=excell,style = "Table Grid")
for i in range(excelh):
    tr = table.rows[i]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), "450")
    trPr.append(trHeight)  # 表格高度设置        
for j in range(excelh):
    array=data[j-1]
    for i in range(excell):
        x_cells = table.rows[j-1].cells
        p = x_cells[i].paragraphs[0]
        run = p.add_run(array[i])
#全文更改字体
from docx.oxml.ns import qn
document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#保存文件
document.save('1.docx')
#页眉页脚
import win32com
from win32com import client
# 打开新的文件 
w = client.Dispatch('Word.Application')  
w.Visible = 0 
w.DisplayAlerts = 0
'''
doc = w.Documents.Open('d:\\biaozhun\\biaozhun.docx') 
w.ActiveDocument.Sections[0].Headers[0].Range.Copy()
wc = client.constants 
doc.Close() 
doc2= w.Documents.Open(path) 
w.ActiveDocument.Sections[0].Headers[0].Range.Paste()
w.ActiveDocument.SaveAs(path)
doc2.Close()
doc3 = w.Documents.Open('d:\\biaozhun\\biaozhun.docx')
w.ActiveDocument.Sections[0].Footers[0].Range.Copy()
doc3.Close()
doc4= w.Documents.Open(path)
w.ActiveDocument.Sections[0].Footers[0].Range.Paste()
w.ActiveDocument.SaveAs(path)
'''
#设置页眉文字，如果要设置页脚值需要把SeekView由9改为10就可以了。。。
doc= w.Documents.Open(path)
w.ActiveWindow.ActivePane.View.SeekView = 9 #9 - 页眉； 10 - 页脚
w.Selection.ParagraphFormat.Alignment = 0
w.Selection.Text = '同济大学'
w.ActiveWindow.ActivePane.View.SeekView = 0 # 释放焦点，返回主文档
w.ActiveWindow.ActivePane.View.SeekView = 10 #9 - 页眉； 10 - 页脚
w.Selection.ParagraphFormat.Alignment = 0
w.Selection.Text = '自动生成报告测试'
w.ActiveWindow.ActivePane.View.SeekView = 0 # 释放焦点，返回主文档
# 页眉文字替换
'''
w.ActiveDocument.Sections[0].Headers[0].Range.Find.ClearFormatting()
w.ActiveDocument.Sections[0].Headers[0].Range.Find.Replacement.ClearFormatting()
w.ActiveDocument.Sections[0].Headers[0].Range.Find.Execute('New Header', False, False, False, False, False, True, 1, False, 'new new header', 2)
'''
w.close()





