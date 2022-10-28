# -*- coding: utf-8 -*-
"""
Created on Sat Sep 22 13:46:33 2018

@author: DELL
"""

import win32com.client
import os
import time
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
import xlrd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
class GENERATEWord():
    def __init__(self,filename=None):
        if filename:
            self.filename=filename
            if os.path.exists(self.filename):
                self.doc=Document(self.filename)
            else:
                self.doc = Document()    #创建新的文档
                self.doc.save(self.filename)
        else:
            self.doc=Document()
            self.filename=''
    def add_title0(self,string):
        self.string=string
        self.doc.add_heading(self.string, level=0)
    def add_title1(self,string):
        self.string=string
        self.doc.add_heading(self.string, level=1)
    def add_title2(self,string):
        self.string=string
        self.doc.add_heading(self.string, level=2)
    def add_text(self,string):
        self.string=string
        self.doc.add_paragraph(self.string)
        #self.doc.add_heading('Heading, level 1', level=1)
        #self.p = self.doc.add_paragraph('A plain paragraph having some ')
        #self.p.add_run('bold').bold = True
        #self.p.add_run(' and some ')
        #self.p.add_run('italic.').italic = True
        #测试缩进
    def add_page_break(self):
        self.doc.add_page_break()#增加分页
        
    #段落测试
    def add_prargraph_with_style(self,string,string1=12,string2=12,string3=8,keep_together=1,keep_with_next=1):
        self.string=string
        self.string1=string1
        self.string2=string2
        self.string3=string3
        self.keep_together=keep_together
        self.keep_with_next=keep_with_next
        self.paragraph = self.doc.add_paragraph(self.string)
        self.paragraph_format =self.paragraph.paragraph_format
        self.paragraph_format.space_before = Pt(self.string1)#本段文字与上一段的间距
        self.paragraph_format.space_after = Pt(self.string2)#本段文字与下一段的间距 
        self.paragraph_format.line_spacing = Pt(self.string3)#行间距
        #分页测试
        if self.keep_together:
            self.paragraph_format.keep_together#紧跟上段
        if self.keep_with_next:
            self.paragraph_format.keep_with_next#若本页无法完全显示，另起一页
    def left_retract(self,string,intches=0.3):
        self.string=string
        self.intches=intches
        self.paragraph = self.doc.add_paragraph(self.string)
        self.paragraph_format = self.paragraph.paragraph_format
        self.paragraph_format.left_indent = Inches(self.intches)
        #首行缩进
    def first_retract(self,string,intches=1):
        self.string=string
        self.intches=intches
        self.paragraph = self.doc.add_paragraph(self.string)
        self.paragraph_format1 =self.paragraph.paragraph_format
        self.paragraph_format1.first_line_indent = Inches(self.intches)
    def add_text_with_style(self,string,string1=1,string2=1,string3=1,string4=28,string5=0,string6=0,string7=0):
        self.string=string
        self.string1=string1#居中
        self.string2=string2#加粗
        self.string3=string3#斜体
        self.string4=string4#大小
        self.string5=string5
        self.string6=string6
        self.string7=string7
        self.p = self.doc.add_paragraph()
        if(self.string1):
            self.p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run = self.p.add_run(self.string)
        self.run.font.color.rgb = RGBColor(self.string5, self.string6, self.string7)#颜色
        self.run.font.size = Pt(self.string4)#字号
        if(self.string2):
            self.run.bold = True
        if(self.string3):
            self.run.italic = True
    def change_style_bold(self,string):
        self.string=string
        self.p = self.doc.add_paragraph()
        self.run = self.p.add_run(self.string)        
        self.run.font.bold = True#加粗
    def change_style_italic(self,string):
        self.string=string
        self.p = self.doc.add_paragraph()
        self.run = self.p.add_run(self.string)
        self.run.font.italic = True#斜体
    def change_style_underline(self,string):
        self.string=string
        self.p = self.doc.add_paragraph()
        self.run = self.p.add_run(self.string)    
        self.run.font.underline = True#下划线
    def change_style_color(self,string):#颜色函数还需要完善
        self.string=string
        self.p = self.doc.add_paragraph()
        self.run = self.p.add_run(self.string)     
        self.run.font.color.rgb = RGBColor(0x42, 0x24 , 0xE9)
    def change_style_size(self,string,size):
        self.string=string
        self.size=size
        self.p = self.doc.add_paragraph()
        self.run = self.p.add_run(self.string)     
        self.run.font.size = Pt(self.size)#字号
    def add_excel(self):
        self.x=xlrd.open_workbook(r"C:\Users\DELL\Desktop\自动报告尝试1\docsample\!.xlsx")
        self.y=self.x.sheets()[0]#获得第一个sheet
        self.excelh=self.y.nrows
        self.excell=self.y.ncols
        self.data=[]
        for i in range(0,self.y.nrows):
            self.a=self.y.row_values(i)
            self.data.append(self.a)
        self.table = self.doc.add_table(rows=self.excelh, cols=self.excell,style = "Table Grid")
        for i in range(self.excelh):
            self.tr = self.table.rows[i]._tr
            self.trPr = self.tr.get_or_add_trPr()
            self.trHeight = OxmlElement('w:trHeight')
            self.trHeight.set(qn('w:val'), "450")
            self.trPr.append(self.trHeight)  # 表格高度设置        
        for j in range(self.excelh):
            self.array=self.data[j-1]
            for i in range(self.excell):
                self.x_cells = self.table.rows[j-1].cells
                self.p = self.x_cells[i].paragraphs[0]
                self.run = self.p.add_run(self.array[i])  
    def add_table(self,rows = 2,cols = 3):
        #表格
        self.rows = rows
        self.cols = cols
        self.table = self.doc.add_table(rows=self.rows, cols=self.cols,style = "Table Grid")  # 添加2行3列的表格
        for i in range(self.rows):
            self.tr = self.table.rows[i]._tr
            self.trPr = self.tr.get_or_add_trPr()
            self.trHeight = OxmlElement('w:trHeight')
            self.trHeight.set(qn('w:val'), "450")
            self.trPr.append(self.trHeight)  # 表格高度设置
            # table.autofit = False
            self.col = self.table.columns[1]
            self.col.width = Inches(5)
            arr = [u'序号',u"类型",u"详细描述"]
            self.heading_cells = self.table.rows[0].cells
            for j in range(cols):
                self.p = self.heading_cells[j].paragraphs[0]
                self.run = self.p.add_run(arr[j])
                self.run.font.color.rgb = RGBColor(54, 95, 145)  # 颜色设置，这里是用RGB颜色
                self.run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                self.p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.table.cell(1, 1).text = u'表格文字'
                #self.table.add_row()
    def add_picture(self,path,width,center):
        self.path=path
        self.width=width
        self.center=center
        self.doc.add_picture(self.path,width=Inches(self.width))
        if(self.center==1):
                last_paragraph = self.doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
    def save_doc(self,path):
        self.path=path
        self.doc.save(self.path)
    def change_all_word_style(self):
        self.doc.styles['Normal'].font.name = u'宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')



        
        
        
        
        
        
        
        
        
        
        
        
        